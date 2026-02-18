import re
from collections import Counter
from pathlib import Path
import fitz  # PyMuPDF
from docx import Document


SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


def extract_text(file_path: str) -> str:
    """Extract and clean text from a PDF or DOCX file."""
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(file_path)
    elif ext == ".docx":
        return _extract_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _extract_pdf(pdf_path: str) -> str:
    """Extract text from a PDF using PyMuPDF."""
    doc = fitz.open(pdf_path)
    pages_text = [page.get_text("text") for page in doc]
    doc.close()
    pages_text = remove_repeated_lines(pages_text)
    return clean_text("\n".join(pages_text))


def _extract_docx(docx_path: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    doc = Document(docx_path)
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            paragraphs.append("")
            continue
        # Prefix headings with a pause-friendly label
        if para.style.name.startswith("Heading"):
            paragraphs.append(f"\n{text}\n")
        else:
            paragraphs.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                paragraphs.append(row_text)

    return clean_text("\n".join(paragraphs))


def remove_repeated_lines(pages: list[str]) -> list[str]:
    """
    Detect lines that appear on many pages (headers/footers) and remove them.
    A line appearing on more than 20% of pages is considered a header/footer.
    """
    if len(pages) < 5:
        return pages

    threshold = max(3, len(pages) * 0.20)

    # Count how often each stripped line appears across pages
    line_counts: Counter = Counter()
    for page in pages:
        seen_on_page = set()
        for line in page.splitlines():
            stripped = line.strip()
            if stripped and stripped not in seen_on_page:
                line_counts[stripped] += 1
                seen_on_page.add(stripped)

    repeated = {line for line, count in line_counts.items() if count >= threshold}

    cleaned = []
    for page in pages:
        lines = [
            line for line in page.splitlines()
            if line.strip() not in repeated
        ]
        cleaned.append("\n".join(lines))

    return cleaned


def clean_text(text: str) -> str:
    """Remove noise commonly found in extracted PDF text."""
    lines = text.splitlines()
    cleaned_lines = []

    for line in lines:
        stripped = line.strip()

        # Skip blank lines temporarily (will restore paragraph breaks)
        if not stripped:
            cleaned_lines.append("")
            continue

        # Remove standalone page numbers: digits only, roman numerals, or "Page N"
        if is_page_number(stripped):
            continue

        # Remove very short noise lines (1-2 chars, single symbols)
        if len(stripped) <= 2 and not stripped[-1].isalpha():
            continue

        # Remove lines that are purely decorative (dashes, underscores, dots)
        if re.fullmatch(r'[-_=~.*•·\s]{3,}', stripped):
            continue

        cleaned_lines.append(line)

    text = "\n".join(cleaned_lines)

    # Fix hyphenation at line breaks (e.g., "some-\nword" -> "someword")
    text = re.sub(r'-\n', '', text)

    # Join lines that are not paragraph breaks (single newline = continuation)
    text = re.sub(r'(?<!\n)\n(?!\n)', ' ', text)

    # Collapse 3+ blank lines into a single paragraph break
    text = re.sub(r'\n{3,}', '\n\n', text)

    # Clean up quotation mark noise — normalize smart quotes to plain
    text = text.replace('\u201c', '"').replace('\u201d', '"')
    text = text.replace('\u2018', "'").replace('\u2019', "'")

    # Remove non-printable / control characters (keep newlines)
    text = re.sub(r'[^\x20-\x7E\n]', ' ', text)

    # Collapse multiple spaces
    text = re.sub(r' {2,}', ' ', text)

    # Remove "quoted" standalone lines that are just a label, e.g. "Chapter 1"
    # but keep them if they're followed by real content (handled by chunking)

    return text.strip()


def is_page_number(line: str) -> bool:
    """Return True if the line looks like a page number or section label to skip."""
    # Plain integer
    if re.fullmatch(r'\d{1,4}', line):
        return True
    # Roman numerals (i, ii, iii, iv, v, vi, vii, viii, ix, x, xi...)
    if re.fullmatch(r'[ivxlcdmIVXLCDM]{1,6}', line):
        return True
    # "Page 42" or "- 42 -" or "42 |" patterns
    if re.fullmatch(r'[-–|]?\s*\d{1,4}\s*[-–|]?', line):
        return True
    # "Chapter 1" / "CHAPTER ONE" alone on a line under 40 chars
    if re.fullmatch(r'(chapter|section|part|book)\s+[\w\s]{1,30}', line, re.IGNORECASE):
        return False  # Keep chapter headings — read them once
    return False


def chunk_text(text: str, max_chars: int = 3000) -> list[str]:
    """Split text into chunks at sentence boundaries for TTS processing."""
    sentences = re.split(r'(?<=[.!?])\s+', text)
    chunks = []
    current = ""

    for sentence in sentences:
        if len(current) + len(sentence) + 1 <= max_chars:
            current += (" " if current else "") + sentence
        else:
            if current:
                chunks.append(current.strip())
            if len(sentence) > max_chars:
                for i in range(0, len(sentence), max_chars):
                    chunks.append(sentence[i:i + max_chars])
            else:
                current = sentence

    if current:
        chunks.append(current.strip())

    return [c for c in chunks if c.strip()]
